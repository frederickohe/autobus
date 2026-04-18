"""Document processor for extracting content from various file formats."""

import logging
from typing import Optional, Tuple
from pathlib import Path
import requests

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Extracts text content from various document formats."""
    
    SUPPORTED_FORMATS = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "application/pdf": "pdf",
        "text/plain": "txt",
        "text/markdown": "md",
    }
    
    @staticmethod
    def extract_text_from_url(file_url: str, file_type: str) -> Optional[str]:
        """Download file from URL and extract text content.
        
        Args:
            file_url: URL to the file
            file_type: MIME type of the file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            # Download file from URL
            response = requests.get(file_url, timeout=30)
            response.raise_for_status()
            
            # Determine format
            format_key = file_type.lower()
            file_format = DocumentProcessor.SUPPORTED_FORMATS.get(format_key)
            
            if not file_format:
                logger.warning(f"Unsupported file format: {file_type}")
                return None
            
            # Extract based on format
            if file_format == "docx":
                return DocumentProcessor._extract_from_docx(response.content)
            elif file_format == "pdf":
                return DocumentProcessor._extract_from_pdf(response.content)
            elif file_format in ["txt", "md"]:
                return response.text
            else:
                return None
                
        except requests.exceptions.RequestException as e:
            logger.error(f"Error downloading document from {file_url}: {e}")
            return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_url}: {e}")
            return None
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> Optional[str]:
        """Extract text from DOCX content.
        
        Args:
            content: File content in bytes
            
        Returns:
            Extracted text
        """
        if not DocxDocument:
            logger.error("python-docx not installed. Install it with: pip install python-docx")
            return None
            
        try:
            from io import BytesIO
            doc = DocxDocument(BytesIO(content))
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            return None
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> Optional[str]:
        """Extract text from PDF content.
        
        Args:
            content: File content in bytes
            
        Returns:
            Extracted text
        """
        if not PdfReader:
            logger.error("PyPDF2 not installed. Install it with: pip install PyPDF2")
            return None
            
        try:
            from io import BytesIO
            pdf_reader = PdfReader(BytesIO(content))
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return None
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> list:
        """Split text into overlapping chunks for better retrieval.
        
        Args:
            text: Full document text
            chunk_size: Size of each chunk in characters
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        text = text.strip()
        
        # Split by sentences/paragraphs first for better semantics
        paragraphs = text.split('\n')
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Add paragraph to current chunk
            if current_chunk:
                test_chunk = current_chunk + "\n" + para
            else:
                test_chunk = para
            
            # If chunk exceeds size, save current and start new
            if len(test_chunk) > chunk_size and current_chunk:
                chunks.append(current_chunk)
                # Create overlap
                current_chunk = current_chunk[-overlap:] + "\n" + para if overlap > 0 else para
            else:
                current_chunk = test_chunk
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        return [c for c in chunks if c.strip()]
